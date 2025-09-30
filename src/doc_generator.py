from docx import Document
import os

DATA_DIR = "../data"
OUTPUT_DIR = "../output"

def fill_template(tool_name, tool_data):
    template_path = os.path.join(DATA_DIR, "Tool Webpage Template_2025.docx")
    doc = Document(template_path)

    # Add title
    doc.add_heading(f"Tool Review: {tool_name}", level=1)

    # Add summary
    if "Tool Finder Summary" in tool_data:
        doc.add_heading("Tool Finder Summary", level=2)
        doc.add_paragraph(tool_data["Tool Finder Summary"])

    # Add status/purpose/usage table
    doc.add_heading("Status / Purpose / Usage", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Status'
    hdr_cells[1].text = 'Purpose'
    hdr_cells[2].text = 'Usage'
    row_cells = table.add_row().cells
    row_cells[0].text = tool_data.get("Status", "")
    row_cells[1].text = ", ".join(tool_data.get("Purpose", []))
    row_cells[2].text = ", ".join(tool_data.get("Usage", []))

    # Add cost and support
    doc.add_heading("Cost", level=2)
    doc.add_paragraph(tool_data.get("Cost", ""))
    doc.add_heading("Support", level=2)
    doc.add_paragraph(tool_data.get("Support", ""))

    # Add overview
    if "Overview" in tool_data:
        doc.add_heading("Overview", level=2)
        doc.add_paragraph(tool_data["Overview"])

    # Add functions
    if "Functions" in tool_data:
        doc.add_heading("Functions", level=2)
        funcs = tool_data["Functions"]
        doc.add_paragraph(funcs.get("Features", ""))
        doc.add_paragraph(f"Grade pass back: {funcs.get('Grade pass back', '')}")
        doc.add_paragraph(f"Canvas LTI: {funcs.get('Canvas LTI', '')}")
        doc.add_paragraph(f"Non-LTI: {funcs.get('Non-LTI', '')}")
        doc.add_paragraph(f"Mobile Access: {funcs.get('Mobile Access', '')}")

    # Save output
    output_path = os.path.join(OUTPUT_DIR, f"{tool_name}_Review.docx")
    doc.save(output_path)
    print(f"Saved: {output_path}")
